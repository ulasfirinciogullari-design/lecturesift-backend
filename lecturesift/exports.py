import json
import shutil
from concurrent.futures import ThreadPoolExecutor
from pathlib import Path
from xml.sax.saxutils import escape

from docx import Document
from docx.shared import Inches
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER
from reportlab.lib.pagesizes import A4, landscape
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import mm
from reportlab.lib.utils import ImageReader
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.platypus import Image as ReportImage
from reportlab.platypus import PageBreak, Paragraph, SimpleDocTemplate, Spacer

from .config import ARTIFACT_EXPORT_PARALLELISM


FONT_PATH = Path("/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf")
FONT_BOLD_PATH = Path("/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf")
FONT_NAME = "Helvetica"
FONT_BOLD = "Helvetica-Bold"

if FONT_PATH.exists():
    pdfmetrics.registerFont(TTFont("LectureSift", str(FONT_PATH)))
    FONT_NAME = "LectureSift"
if FONT_BOLD_PATH.exists():
    pdfmetrics.registerFont(TTFont("LectureSift-Bold", str(FONT_BOLD_PATH)))
    FONT_BOLD = "LectureSift-Bold"


def _styles() -> dict:
    base = getSampleStyleSheet()
    return {
        "title": ParagraphStyle(
            "LectureTitle",
            parent=base["Title"],
            fontName=FONT_BOLD,
            fontSize=22,
            leading=27,
            textColor=colors.HexColor("#172554"),
            alignment=TA_CENTER,
            spaceAfter=14,
        ),
        "heading": ParagraphStyle(
            "LectureHeading",
            parent=base["Heading2"],
            fontName=FONT_BOLD,
            fontSize=14,
            leading=18,
            textColor=colors.HexColor("#4338CA"),
            spaceBefore=10,
            spaceAfter=7,
        ),
        "body": ParagraphStyle(
            "LectureBody",
            parent=base["BodyText"],
            fontName=FONT_NAME,
            fontSize=10.5,
            leading=15,
            textColor=colors.HexColor("#1F2937"),
            spaceAfter=6,
        ),
        "small": ParagraphStyle(
            "LectureSmall",
            parent=base["BodyText"],
            fontName=FONT_NAME,
            fontSize=8.5,
            leading=12,
            textColor=colors.HexColor("#64748B"),
        ),
    }


def _paragraph(text: str, style) -> Paragraph:
    return Paragraph(escape(str(text or "")).replace("\n", "<br/>"), style)


def _write_pdf(path: Path, title: str, sections: list[tuple[str, list[str]]]) -> None:
    styles = _styles()
    document = SimpleDocTemplate(
        str(path),
        pagesize=A4,
        rightMargin=18 * mm,
        leftMargin=18 * mm,
        topMargin=17 * mm,
        bottomMargin=17 * mm,
        title=title,
        author="LectureSift",
    )
    story = [_paragraph(title, styles["title"]), _paragraph("LectureSift AI Study Pack", styles["small"]), Spacer(1, 8)]
    for heading, paragraphs in sections:
        if heading:
            story.append(_paragraph(heading, styles["heading"]))
        for value in paragraphs:
            story.append(_paragraph(value, styles["body"]))
    document.build(story)


def _write_docx(path: Path, title: str, sections: list[tuple[str, list[str]]]) -> None:
    document = Document()
    document.add_heading(title, 0)
    document.add_paragraph("LectureSift AI Study Pack")
    for heading, paragraphs in sections:
        if heading:
            document.add_heading(heading, level=1)
        for value in paragraphs:
            for paragraph in str(value or "").split("\n"):
                document.add_paragraph(paragraph)
    document.save(path)


def _notes_text(pack: dict) -> str:
    blocks: list[str] = []
    if pack.get("key_points"):
        blocks.append("ÖNEMLİ NOKTALAR\n" + "\n".join(f"• {item}" for item in pack["key_points"]))
    if pack.get("important_terms"):
        blocks.append(
            "KAVRAMLAR VE TANIMLAR\n"
            + "\n".join(f"• {item.get('term', '')}: {item.get('definition', '')}" for item in pack["important_terms"])
        )
    for note in pack.get("notes", []):
        value = f"{note.get('heading', '')}\n{note.get('content', '')}"
        bullets = note.get("bullets") or []
        if bullets:
            value += "\n" + "\n".join(f"• {bullet}" for bullet in bullets)
        blocks.append(value.strip())
    if pack.get("exam_focus"):
        blocks.append("SINAV ODAKLI NOKTALAR\n" + "\n".join(f"• {item}" for item in pack["exam_focus"]))
    return "\n\n".join(blocks).strip()


def _quiz_text(quiz: list[dict]) -> str:
    blocks = []
    for index, item in enumerate(quiz, 1):
        options = "\n".join(f"  {chr(65 + option_index)}. {option}" for option_index, option in enumerate(item.get("options", [])))
        answer_index = int(item.get("answer_index", 0))
        blocks.append(
            f"{index}. {item.get('question', '')}\n{options}\n"
            f"Doğru cevap: {chr(65 + answer_index)}\nAçıklama: {item.get('explanation', '')}"
        )
    return "\n\n".join(blocks)


def _flashcards_text(cards: list[dict]) -> str:
    return "\n\n".join(
        f"{index}. Soru: {item.get('front', '')}\nCevap: {item.get('back', '')}"
        for index, item in enumerate(cards, 1)
    )


def _write_slides_pdf(path: Path, title: str, slides: list[dict], slides_dir: Path) -> None:
    styles = _styles()
    document = SimpleDocTemplate(
        str(path),
        pagesize=landscape(A4),
        rightMargin=15 * mm,
        leftMargin=15 * mm,
        topMargin=12 * mm,
        bottomMargin=12 * mm,
        title=f"{title} - Slaytlar",
        author="LectureSift",
    )
    story = []
    max_width, max_height = 260 * mm, 155 * mm
    for index, slide in enumerate(slides):
        image_path = slides_dir / slide["file"]
        if not image_path.exists():
            continue
        width, height = ImageReader(str(image_path)).getSize()
        scale = min(max_width / width, max_height / height)
        story.append(_paragraph(f"{title} — {slide.get('timestamp', '')}", styles["heading"]))
        story.append(ReportImage(str(image_path), width=width * scale, height=height * scale))
        if index < len(slides) - 1:
            story.append(PageBreak())
    if not story:
        story.append(_paragraph("Slayt bulunamadı.", styles["body"]))
    document.build(story)


def _write_slides_docx(path: Path, title: str, slides: list[dict], slides_dir: Path) -> None:
    document = Document()
    document.add_heading(f"{title} - Slaytlar", 0)
    for index, slide in enumerate(slides):
        image_path = slides_dir / slide["file"]
        if not image_path.exists():
            continue
        document.add_heading(slide.get("timestamp", ""), level=1)
        document.add_picture(str(image_path), width=Inches(6.5))
        if index < len(slides) - 1:
            document.add_page_break()
    document.save(path)


def _artifact(path: Path, label: str) -> dict:
    return {
        "file": path.name,
        "label": label,
        "format": path.suffix.removeprefix(".").upper(),
        "size_bytes": path.stat().st_size,
    }


def _save_result(job_dir: Path, result: dict, artifacts: list[dict]) -> None:
    complete_result = {**result, "artifacts": artifacts}
    (job_dir / "result.json").write_text(json.dumps(complete_result, ensure_ascii=False, indent=2), encoding="utf-8")


def _package_audio_artifact(job_dir: Path, package_dir: Path, source: Path | None) -> dict | None:
    """Copy one worker-generated audio file into the private download package.

    The caller always supplies a server-generated path.  Still constrain it to
    this job and reject links so a future caller cannot turn the exporter into
    an arbitrary-file read.  A fixed archive name also prevents source
    filenames from becoming ZIP member paths.
    """
    if source is None:
        return None
    source = Path(source)
    if source.is_symlink() or not source.is_file():
        raise ValueError("Study-pack audio must be a regular file.")
    try:
        resolved_job = job_dir.resolve(strict=True)
        resolved_source = source.resolve(strict=True)
        resolved_source.relative_to(resolved_job)
    except (OSError, ValueError) as exc:
        raise ValueError("Study-pack audio escaped its job directory.") from exc

    destination = package_dir / "LectureSift_Ders_Sesi.mp3"
    if resolved_source != destination.resolve(strict=False):
        destination.unlink(missing_ok=True)
        shutil.copyfile(resolved_source, destination)
    if destination.is_symlink() or not destination.is_file() or destination.stat().st_size <= 0:
        raise ValueError("Study-pack audio is empty or invalid.")
    return _artifact(destination, "Ders Sesi (MP3)")


def _timestamped_transcript(result: dict) -> str:
    segments = result.get("transcript_segments") or []
    if not segments:
        return result.get("transcript_original", "") or "Videoda ses parçası bulunamadı."
    rows = []
    for segment in segments:
        speaker = str(segment.get("speaker") or "").strip()
        label = f" {speaker}" if speaker else ""
        rows.append(f"[{segment.get('timestamp', '00:00:00')}]{label}  {segment.get('text', '')}".strip())
    return "\n\n".join(rows)


def build_artifacts(
    job_dir: Path,
    result: dict,
    slides_dir: Path,
    *,
    audio_source: Path | None = None,
    notes_stem: str = "Ders_Notlari",
    notes_label: str = "Ders Notları",
    archive_stem: str = "LectureSift_Study_Pack_V4",
) -> tuple[list[dict], Path]:
    package_dir = job_dir / "package"
    package_dir.mkdir(parents=True, exist_ok=True)
    options = result.get("options", {})
    configured_formats = options.get("output_formats")
    formats = set(["pdf"] if configured_formats is None else configured_formats)
    title = result.get("title") or "LectureSift Ders Paketi"
    original = _timestamped_transcript(result)
    translated = result.get("transcript_translated", "")
    documents = []
    if options.get("include_summary", True):
        notes_text = _notes_text(result)
        documents.extend(
            [
                ("Ozet", "Özet", f"{title} - Özet", [("Özet", [result.get("summary", "")])], result.get("summary", "")),
                (notes_stem, notes_label, f"{title} - Ders Notları", [("Ders Notları", [notes_text])], notes_text),
            ]
        )
    if options.get("include_transcript", True):
        documents.append(("Transkript_Orijinal", "Orijinal Transkript", f"{title} - Orijinal Transkript", [("Transkript", [original])], original))
        if translated:
            documents.append(("Transkript_Ceviri", "Çevrilmiş Transkript", f"{title} - Çevrilmiş Transkript", [("Transkript", [translated])], translated))
    if result.get("quiz"):
        documents.append(("Quiz", "Quiz", f"{title} - Quiz", [("Sorular ve Yanıtlar", [_quiz_text(result["quiz"])])], _quiz_text(result["quiz"])))
    if result.get("flashcards"):
        documents.append(("Flashcards", "Bilgi Kartları", f"{title} - Bilgi Kartları", [("Bilgi Kartları", [_flashcards_text(result["flashcards"])])], _flashcards_text(result["flashcards"])))

    export_jobs: list[tuple[Path, str, object, tuple]] = []
    for stem, label, document_title, sections, plain_text in documents:
        if "pdf" in formats:
            path = package_dir / f"{stem}.pdf"
            export_jobs.append((path, f"{label} (PDF)", _write_pdf, (path, document_title, sections)))
        if "docx" in formats:
            path = package_dir / f"{stem}.docx"
            export_jobs.append((path, f"{label} (Word)", _write_docx, (path, document_title, sections)))
        if "txt" in formats:
            path = package_dir / f"{stem}.txt"
            export_jobs.append((path, f"{label} (TXT)", Path.write_text, (path, plain_text, "utf-8")))

    slides = result.get("slides", []) if options.get("include_slides", True) else []
    if slides:
        if "pdf" in formats:
            path = package_dir / "Slaytlar.pdf"
            export_jobs.append((path, "Slaytlar (PDF)", _write_slides_pdf, (path, title, slides, slides_dir)))
        if "docx" in formats:
            path = package_dir / "Slaytlar.docx"
            export_jobs.append((path, "Slaytlar (Word)", _write_slides_docx, (path, title, slides, slides_dir)))
        if "txt" in formats:
            path = package_dir / "Slaytlar.txt"
            plain_text = "\n".join(f"{item.get('timestamp', '')} — {item.get('file', '')}" for item in slides)
            export_jobs.append((path, "Slayt Listesi (TXT)", Path.write_text, (path, plain_text, "utf-8")))

    def export(item: tuple[Path, str, object, tuple]) -> dict:
        path, label, writer, arguments = item
        writer(*arguments)
        return _artifact(path, label)

    workers = min(ARTIFACT_EXPORT_PARALLELISM, len(export_jobs))
    if workers <= 1:
        artifacts = [export(item) for item in export_jobs]
    else:
        with ThreadPoolExecutor(
            max_workers=workers,
            thread_name_prefix="lecturesift-export",
        ) as executor:
            artifacts = list(executor.map(export, export_jobs))

    audio_artifact = _package_audio_artifact(job_dir, package_dir, audio_source)
    if audio_artifact is not None:
        artifacts.append(audio_artifact)

    _save_result(job_dir, result, artifacts)
    zip_base = job_dir / archive_stem
    shutil.make_archive(str(zip_base), "zip", root_dir=package_dir)
    return artifacts, zip_base.with_suffix(".zip")


def build_binary_artifact(job_dir: Path, result: dict, source: Path, filename: str, label: str) -> tuple[list[dict], Path]:
    package_dir = job_dir / "package"
    package_dir.mkdir(parents=True, exist_ok=True)
    destination = package_dir / filename
    shutil.copy2(source, destination)
    artifacts = [_artifact(destination, label)]
    _save_result(job_dir, result, artifacts)
    zip_base = job_dir / "LectureSift_Download"
    shutil.make_archive(str(zip_base), "zip", root_dir=package_dir)
    return artifacts, zip_base.with_suffix(".zip")
